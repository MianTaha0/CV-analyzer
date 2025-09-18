from flask import Flask, request, jsonify, send_file
import os, json, re, io, tempfile
from werkzeug.utils import secure_filename
from pdfminer.high_level import extract_text as extract_pdf_text
import docx
from pymongo import MongoClient
from bson import ObjectId
from flask_cors import CORS
from openai import OpenAI
import numpy as np
from dotenv import load_dotenv
 # <-- replace with your ffmpeg path




load_dotenv()

# ==========================
# CONFIG
# ==========================
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# OpenAI Client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# MongoDB Init
mongo_client = MongoClient("mongodb://localhost:27017/")
db = mongo_client["cv_analyzer"]
candidates_collection = db["candidates"]
candidates_collection.create_index("email", unique=True, sparse=True)

# Flask App
app = Flask(__name__)
CORS(app)


# ==========================
# HELPERS
# ==========================
def safe_json_parse(text, fallback=None):
    try:
        match = re.search(r"\{.*\}", text, re.S)
        if match:
            return json.loads(match.group())
    except Exception as e:
        print("JSON parse error:", e)
    return fallback


def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_pdf_text(file_path)
    elif ext in [".doc", ".docx"]:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        raise ValueError("Unsupported file type")


# ==========================
# ROUTES
# ==========================
@app.route("/upload", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "No filename"}), 400

    filepath = os.path.join(UPLOAD_DIR, secure_filename(file.filename))
    file.save(filepath)

    # Extract text
    try:
        text = extract_text(filepath)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    # Parse with GPT
    prompt = f"""
    Extract structured details from this resume.
    Return ONLY JSON with this structure:
    {{
      "name": "",
      "email": "",
      "phone": "",
      "skills": [],
      "experience": [],
      "education" : [],
      "location" : "",
      "total_experience": {{"years": 0, "formatted": ""}},
      "role": ""
    }}
    Resume text:
    {text}
    """
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0
        )
        parsed = safe_json_parse(resp.choices[0].message.content, fallback={})
        usage = resp.usage.to_dict() if resp.usage else {}
    except Exception as e:
        return jsonify({"error": f"AI parsing failed: {e}"}), 500

    if not parsed.get("email"):
        return jsonify({"error": "Could not extract candidate info"}), 500

    # Generate embedding
    embedding_input = f"{parsed.get('name','')} {parsed.get('skills')} {parsed.get('experience')} {text}"
    embedding = client.embeddings.create(
        model="text-embedding-3-small",
        input=embedding_input
    ).data[0].embedding

    # Save in DB
    candidate_doc = {
        **parsed,
        "raw_text": text,
        "embedding": embedding
    }
    candidates_collection.update_one(
        {"email": parsed["email"]},
        {"$set": candidate_doc},
        upsert=True
    )

    return jsonify({
        "message": "Resume uploaded",
        "candidate": parsed,
        "tokens_used": usage
    })


@app.route("/search", methods=["GET"])
def search_candidates():
    q = request.args.get("q", "").lower()
    if not q:
        return jsonify({"error": "Query required"}), 400

    candidates = list(candidates_collection.find({}))
    results = []
    for c in candidates:
        haystack = f"{c.get('name','')} {c.get('skills')} {c.get('experience')} {c.get('role','')}".lower()
        if q in haystack:
            results.append({
                "id": str(c["_id"]),
                "name": c.get("name"),
                "email": c.get("email"),
                "skills": c.get("skills"),
                "role": c.get("role")
            })

    return jsonify({"query": q, "results": results})


@app.route("/semantic-search", methods=["GET"])
def semantic_search():
    query = request.args.get("q", "")
    if not query:
        return jsonify({"error": "Query is required"}), 400

    q_emb = client.embeddings.create(
        model="text-embedding-3-small",
        input=query
    ).data[0].embedding

    candidates = list(candidates_collection.find({"embedding": {"$exists": True}}))

    def cosine(a, b):
        a, b = np.array(a), np.array(b)
        return float(np.dot(a, b) / ((np.linalg.norm(a) * np.linalg.norm(b)) + 1e-9))

    ranked = []
    for c in candidates:
        sim = cosine(q_emb, c["embedding"])
        ranked.append((sim, c))

    ranked.sort(reverse=True, key=lambda x: x[0])

    results = []
    for sim, c in ranked[:5]:
        results.append({
            "id": str(c["_id"]),
            "name": c.get("name"),
            "email": c.get("email"),
            "skills": c.get("skills"),
            "role": c.get("role"),
            "similarity": round(sim, 3)
        })

    return jsonify({"query": query, "results": results})


@app.route("/ai-job-match", methods=["POST", "GET"])
def ai_job_match():
    if request.method == "POST":
        data = request.get_json() or {}
        requirement = data.get("requirement", "").lower()
    else:  # GET request with ?q=
        requirement = request.args.get("q", "").lower()

    if not requirement:
        return jsonify({"error": "Requirement text is required"}), 400

    # Fetch all candidates
    candidates = list(candidates_collection.find({}))

    # Pre-filter
    filtered_candidates = []
    for c in candidates:
        skills_text = " ".join(c.get("skills", [])).lower()
        role_text = (c.get("role") or "").lower()
        if any(word in requirement for word in skills_text.split()) or any(word in requirement for word in role_text.split()):
            filtered_candidates.append(c)

    if not filtered_candidates:
        return jsonify({
            "requirement": requirement,
            "results": [],
            "message": "No candidate with required skills found"
        })

    # Summaries
    candidate_summaries = []
    for c in filtered_candidates:
        candidate_summaries.append({
            "id": str(c["_id"]),
            "name": c.get("name"),
            "email": c.get("email"),
            "skills": c.get("skills"),
            "total_experience": c.get("total_experience"),
            "experience": c.get("experience"),
            "role": c.get("role")
        })

    system_prompt = """
    You are an expert technical recruiter.
    Compare the given job requirement with the list of candidates.
    For each candidate, give a match score from 0-100 and explain briefly why.
    Only return JSON in this format:
    {
      "results": [
        {
          "id": "",
          "name": "",
          "email": "",
          "score": 0,
          "reason": ""
        }
      ]
    }
    """

    user_prompt = f"""
    Job Requirement:
    {requirement}

    Candidates:
    {json.dumps(candidate_summaries, indent=2)}
    """

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=1500,
            temperature=0
        )
        match_results = safe_json_parse(resp.choices[0].message.content, fallback={"results": []})
        usage = resp.usage.to_dict() if resp.usage else {}
    except Exception as e:
        return jsonify({"error": f"AI matching failed: {e}"}), 500

    results = match_results.get("results", [])
    if not results:
        return jsonify({
            "requirement": requirement,
            "results": [],
            "message": "No candidate matched after AI scoring",
            "tokens_used": usage
        })

    results.sort(key=lambda x: x.get("score", 0), reverse=True)

    return jsonify({
        "requirement": requirement,
        "results": results[:5],
        "tokens_used": usage
    })


# ==========================
# VOICE ROUTE (FIXED)
# ==========================
# ==========================
# VOICE ROUTE (FIXED with pydub conversion)
# ==========================
@app.route("/voice-job-match", methods=["POST"])
def voice_job_match():
    try:
        if "file" not in request.files:
            return jsonify({"error": "No audio file uploaded"}), 400

        audio_file = request.files["file"]
        if not audio_file.filename:
            return jsonify({"error": "No filename"}), 400

        ext = os.path.splitext(audio_file.filename)[1].lower().replace('.', '')
        allowed = ['flac', 'm4a', 'mp3', 'mp4', 'mpeg', 'mpga', 'oga', 'ogg', 'wav', 'webm']
        if ext not in allowed:
            return jsonify({"error": f"Unsupported file type: {ext}"}), 400

        # Save uploaded audio temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmpfile:
            audio_file.save(tmpfile.name)
            audio_path = tmpfile.name

        # 1️⃣ Speech-to-Text (Whisper)
        with open(audio_path, "rb") as f:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=f
            )
        requirement = transcript.text.strip()
        print("DEBUG: Job requirement:", requirement)

        # 2️⃣ Fetch all candidates from MongoDB
        candidates = list(candidates_collection.find({}))
        if not candidates:
            return jsonify({"error": "No candidates saved"}), 400

        candidate_summaries = []
        for c in candidates:
            candidate_summaries.append({
                "name": c.get("name"),
                "email": c.get("email"),
                "skills": c.get("skills"),
                "role": c.get("role"),
                "experience": c.get("experience"),
                "total_experience": c.get("total_experience")
            })

        # 3️⃣ Generate GPT prompt to pick best candidate
        system_prompt = """
        You are a senior recruiter.
        Compare the job requirement with the list of candidates.
        Pick the best candidate, score them 0-100, and explain why they are best.
        Keep it concise, professional, and human-readable.
        """

        user_prompt = f"""
        Job Requirement:
        {requirement}

        Candidates:
        {json.dumps(candidate_summaries, indent=2)}
        """

        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=500,
            temperature=0
        )
        explanation = resp.choices[0].message.content.strip()
        print("DEBUG: Explanation:", explanation)

        # 4️⃣ Text-to-Speech (TTS)
        speech = client.audio.speech.create(
            model="gpt-4o-mini-tts",
            voice="alloy",
            input=explanation
        )
        audio_bytes = speech.read()
        audio_stream = io.BytesIO(audio_bytes)
        audio_stream.seek(0)

        return send_file(
            audio_stream,
            mimetype="audio/mpeg",
            as_attachment=True,
            download_name="best_candidate.mp3"
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Server failed: {str(e)}"}), 500


# ==========================
# RUN
# ==========================
if __name__ == "__main__":
    app.run(debug=True, port=5000)














'''from flask import Flask, request, jsonify
import os, json, re
from werkzeug.utils import secure_filename
from pdfminer.high_level import extract_text as extract_pdf_text
import docx
from pymongo import MongoClient
from bson import ObjectId
from flask_cors import CORS
from openai import OpenAI
import numpy as np
from dotenv import load_dotenv

load_dotenv()

# ==========================
# CONFIG
# ==========================
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# OpenAI Client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# MongoDB Init
mongo_client = MongoClient("mongodb://localhost:27017/")
db = mongo_client["cv_analyzer"]
candidates_collection = db["candidates"]
candidates_collection.create_index("email", unique=True, sparse=True)

# Flask App
app = Flask(__name__)
CORS(app)


# ==========================
# HELPERS
# ==========================
def safe_json_parse(text, fallback=None):
    try:
        match = re.search(r"\{.*\}", text, re.S)
        if match:
            return json.loads(match.group())
    except Exception as e:
        print("JSON parse error:", e)
    return fallback


def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_pdf_text(file_path)
    elif ext in [".doc", ".docx"]:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        raise ValueError("Unsupported file type")


# ==========================
# ROUTES
# ==========================
@app.route("/upload", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "No filename"}), 400

    filepath = os.path.join(UPLOAD_DIR, secure_filename(file.filename))
    file.save(filepath)

    # Extract text
    try:
        text = extract_text(filepath)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    # Parse with GPT
    prompt = f"""
    Extract structured details from this resume.
    Return ONLY JSON with this structure:
    {{
      "name": "",
      "email": "",
      "phone": "",
      "skills": [],
      "experience": [],
      "education" : [],
      "location" : "",
      "total_experience": {{"years": 0, "formatted": ""}},
      "role": ""
    }}
    Resume text:
    {text}
    """
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0
        )
        parsed = safe_json_parse(resp.choices[0].message.content, fallback={})
        usage = resp.usage.to_dict() if resp.usage else {}
    except Exception as e:
        return jsonify({"error": f"AI parsing failed: {e}"}), 500

    if not parsed.get("email"):
        return jsonify({"error": "Could not extract candidate info"}), 500

    # Generate embedding
    embedding_input = f"{parsed.get('name','')} {parsed.get('skills')} {parsed.get('experience')} {text}"
    embedding = client.embeddings.create(
        model="text-embedding-3-small",
        input=embedding_input
    ).data[0].embedding

    # Save in DB
    candidate_doc = {
        **parsed,
        "raw_text": text,
        "embedding": embedding
    }
    candidates_collection.update_one(
        {"email": parsed["email"]},
        {"$set": candidate_doc},
        upsert=True
    )

    return jsonify({
        "message": "Resume uploaded",
        "candidate": parsed,
        "tokens_used": usage
    })


@app.route("/search", methods=["GET"])
def search_candidates():
    q = request.args.get("q", "").lower()
    if not q:
        return jsonify({"error": "Query required"}), 400

    candidates = list(candidates_collection.find({}))
    results = []
    for c in candidates:
        haystack = f"{c.get('name','')} {c.get('skills')} {c.get('experience')} {c.get('role','')}".lower()
        if q in haystack:
            results.append({
                "id": str(c["_id"]),
                "name": c.get("name"),
                "email": c.get("email"),
                "skills": c.get("skills"),
                "role": c.get("role")
            })

    return jsonify({"query": q, "results": results})


@app.route("/semantic-search", methods=["GET"])
def semantic_search():
    query = request.args.get("q", "")
    if not query:
        return jsonify({"error": "Query is required"}), 400

    q_emb = client.embeddings.create(
        model="text-embedding-3-small",
        input=query
    ).data[0].embedding

    candidates = list(candidates_collection.find({"embedding": {"$exists": True}}))

    def cosine(a, b):
        a, b = np.array(a), np.array(b)
        return float(np.dot(a, b) / ((np.linalg.norm(a) * np.linalg.norm(b)) + 1e-9))

    ranked = []
    for c in candidates:
        sim = cosine(q_emb, c["embedding"])
        ranked.append((sim, c))

    ranked.sort(reverse=True, key=lambda x: x[0])

    results = []
    for sim, c in ranked[:5]:
        results.append({
            "id": str(c["_id"]),
            "name": c.get("name"),
            "email": c.get("email"),
            "skills": c.get("skills"),
            "role": c.get("role"),
            "similarity": round(sim, 3)
        })

    return jsonify({"query": query, "results": results})


@app.route("/ai-job-match", methods=["POST", "GET"])
def ai_job_match():
    if request.method == "POST":
        data = request.get_json() or {}
        requirement = data.get("requirement", "").lower()
    else:  # GET request with ?q=
        requirement = request.args.get("q", "").lower()

    if not requirement:
        return jsonify({"error": "Requirement text is required"}), 400

    # Fetch all candidates
    candidates = list(candidates_collection.find({}))

    # Pre-filter: only candidates whose skills/role appear in the requirement
    filtered_candidates = []
    for c in candidates:
        skills_text = " ".join(c.get("skills", [])).lower()
        role_text = (c.get("role") or "").lower()
        if any(word in requirement for word in skills_text.split()) or any(word in requirement for word in role_text.split()):
            filtered_candidates.append(c)

    if not filtered_candidates:
        return jsonify({
            "requirement": requirement,
            "results": [],
            "message": "No candidate with required skills found"
        })

    # Prepare summary for AI
    candidate_summaries = []
    for c in filtered_candidates:
        candidate_summaries.append({
            "id": str(c["_id"]),
            "name": c.get("name"),
            "email": c.get("email"),
            "skills": c.get("skills"),
            "total_experience": c.get("total_experience"),
            "experience": c.get("experience"),
            "role": c.get("role")
        })

    system_prompt = """
    You are an expert technical recruiter.
    Compare the given job requirement with the list of candidates.
    For each candidate, give a match score from 0-100 and explain briefly why.
    Only return JSON in this format:
    {
      "results": [
        {
          "id": "",
          "name": "",
          "email": "",
          "score": 0,
          "reason": ""
        }
      ]
    }
    """

    user_prompt = f"""
    Job Requirement:
    {requirement}

    Candidates:
    {json.dumps(candidate_summaries, indent=2)}
    """

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=1500,
            temperature=0
        )
        match_results = safe_json_parse(resp.choices[0].message.content, fallback={"results": []})
        usage = resp.usage.to_dict() if resp.usage else {}
    except Exception as e:
        return jsonify({"error": f"AI matching failed: {e}"}), 500

    results = match_results.get("results", [])
    if not results:
        return jsonify({
            "requirement": requirement,
            "results": [],
            "message": "No candidate matched after AI scoring",
            "tokens_used": usage
        })

    # Sort by score
    results.sort(key=lambda x: x.get("score", 0), reverse=True)

    return jsonify({
        "requirement": requirement,
        "results": results[:5],
        "tokens_used": usage
    })


# ==========================
# RUN
# ==========================
if __name__ == "__main__":
    app.run(debug=True, port=5000)'''







'''from flask import Flask, request, jsonify
import os, json, re, uuid
from werkzeug.utils import secure_filename
from pdfminer.high_level import extract_text as extract_pdf_text
from dotenv import load_dotenv
import docx
from pymongo import MongoClient
from bson import ObjectId
from flask_cors import CORS
from datetime import datetime
from openai import OpenAI

# ==========================
# CONFIG
# ==========================
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
ALLOWED_EXT = {"pdf", "docx", "txt"}

# Load API key
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
if not os.getenv("OPENAI_API_KEY"):
    raise RuntimeError("OPENAI_API_KEY not set in .env")

app = Flask(__name__)
CORS(app)

# ==========================
# MONGODB INIT
# ==========================
client_mongo = MongoClient("mongodb://localhost:27017/")
db = client_mongo["cv_analyzer"]
candidates_collection = db["candidates"]
candidates_collection.create_index("email", unique=True, sparse=True)

# ==========================
# FILE TEXT EXTRACTION
# ==========================
def extract_text(path):
    ext = path.split(".")[-1].lower()
    if ext == "pdf":
        return extract_pdf_text(path)
    elif ext == "docx":
        doc = docx.Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    elif ext == "txt":
        return open(path, "r", encoding="utf-8", errors="ignore").read()
    return ""

# ==========================
# HELPER: JSON SAFE PARSING
# ==========================
def safe_json_parse(text, fallback=None):
    m = re.search(r"\{(?:.|\n)*\}", text)
    if not m:
        return fallback or {}
    try:
        return json.loads(m.group())
    except Exception as e:
        print("JSON parse error:", e, "RAW:", m.group()[:500])
        return fallback or {}

# ==========================
# OPENAI PARSING
# ==========================
def parse_resume(resume_text):
    prompt = f"""
    Extract the following details from the resume and return only valid JSON.
    If a field is missing, return null or an empty list. Dates for experience
    should include "start" and "end" in format 'MMM YYYY' if possible.

    {{
      "name": "",
      "email": "",
      "phone": "",
      "location": "",
      "education": [],
      "experience": [
         {{"company": "", "role": "", "start": "", "end": ""}}
      ],
      "skills": []
    }}

    Resume:
    {resume_text[:20000]}
    """

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0
        )
        text = resp.choices[0].message.content
        parsed = safe_json_parse(text, fallback={
            "name": None, "email": None, "phone": None,
            "skills": [], "experience": [], "education": [], "location": None
        })
        return parsed, resp.usage.model_dump()   # ✅ return parsed + usage
    except Exception as e:
        print("OpenAI API error:", e)
        return {
            "name": None, "email": None, "phone": None,
            "skills": [], "experience": [], "education": [], "location": None
        }, {}

def semantic_analysis(resume_text):
    prompt = f"""
    Analyze the resume and return only JSON:
    {{
      "summary": "Short 3-4 line professional summary",
      "domain": "IT / Finance / Marketing / Healthcare / Other",
      "role_category": "Developer / Designer / Manager / Analyst / Other",
      "strengths": ["..."],
      "weaknesses": ["..."]
    }}

    Resume:
    {resume_text[:5000]}
    """

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
            temperature=0
        )
        text = resp.choices[0].message.content
        parsed = safe_json_parse(text, fallback={
            "summary": None, "domain": None, "role_category": None,
            "strengths": [], "weaknesses": []
        })
        return parsed, resp.usage.model_dump()   # ✅ return parsed + usage
    except Exception as e:
        print("Semantic API error:", e)
        return {
            "summary": None, "domain": None, "role_category": None,
            "strengths": [], "weaknesses": []
        }, {}

# ==========================
# EXPERIENCE CALCULATION
# ==========================
def calculate_total_experience(experiences):
    total_months = 0
    for exp in experiences:
        start_str = exp.get("start")
        end_str = exp.get("end")
        if not start_str:
            continue

        try:
            start = datetime.strptime(start_str, "%b %Y")
        except:
            try:
                start = datetime.strptime(start_str, "%Y")
            except:
                continue

        if end_str:
            try:
                end = datetime.strptime(end_str, "%b %Y")
            except:
                try:
                    end = datetime.strptime(end_str, "%Y")
                except:
                    end = datetime.today()
        else:
            end = datetime.today()

        total_months += (end.year - start.year) * 12 + (end.month - start.month)

    total_years = round(total_months / 12, 2)
    return {
        "years": total_years,
        "formatted": f"{int(total_months/12)} years {total_months%12} months"
    }

# ==========================
# ROUTES
# ==========================
@app.route("/upload", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file"}), 400

    file = request.files["file"]
    filename = secure_filename(file.filename)
    if not filename or filename.split(".")[-1].lower() not in ALLOWED_EXT:
        return jsonify({"error": "Invalid file type"}), 400

    unique_name = f"{uuid.uuid4().hex}_{filename}"
    path = os.path.join(UPLOAD_DIR, unique_name)
    file.save(path)

    text = extract_text(path)
    print("=== Extracted Resume Text ===")
    print(text[:1000])

    parsed, usage_resume = parse_resume(text)
    semantic, usage_semantic = semantic_analysis(text)
    parsed["total_experience"] = calculate_total_experience(parsed.get("experience", []))

    # Duplicate check
    existing = None
    if parsed.get("email"):
        existing = candidates_collection.find_one({"email": parsed.get("email")})
    if not existing:
        existing = candidates_collection.find_one({"raw_text": text})

    if existing:
        return jsonify({
            "ok": False,
            "message": "This CV is already in the database",
            "id": str(existing["_id"]),
            "parsed": {
                "name": existing.get("name"),
                "email": existing.get("email"),
                "phone": existing.get("phone"),
                "skills": existing.get("skills"),
                "experience": existing.get("experience"),
                "education": existing.get("education"),
                "location": existing.get("location"),
                "total_experience": existing.get("total_experience"),
                "semantic": existing.get("semantic")
            }
        }), 409

    candidate_doc = {
        "name": parsed.get("name"),
        "email": parsed.get("email"),
        "phone": parsed.get("phone"),
        "skills": parsed.get("skills") or [],
        "experience": parsed.get("experience") or [],
        "education": parsed.get("education") or [],
        "location": parsed.get("location"),
        "total_experience": parsed.get("total_experience"),
        "semantic": semantic,
        "raw_text": text
    }
    inserted = candidates_collection.insert_one(candidate_doc)

    return jsonify({
        "ok": True,
        "id": str(inserted.inserted_id),
        "parsed": parsed,
        "semantic": semantic,
        "token_usage": {
            "resume_parsing": usage_resume,
            "semantic_analysis": usage_semantic
        }
    })

@app.route("/search")
def search():
    skill_query = request.args.get("skill", "")
    name_query = request.args.get("name", "")
    email_query = request.args.get("email", "")

    query_parts = []
    if skill_query:
        skills = [s.strip() for s in skill_query.split(",")]
        query_parts.append({"$or": [{"skills": {"$regex": s, "$options": "i"}} for s in skills]})
    if name_query:
        query_parts.append({"name": {"$regex": re.escape(name_query), "$options": "i"}})
    if email_query:
        query_parts.append({"email": {"$regex": re.escape(email_query), "$options": "i"}})

    query = {"$and": query_parts} if query_parts else {}
    cursor = candidates_collection.find(query)

    results = []
    for doc in cursor:
        results.append({
            "id": str(doc["_id"]),
            "name": doc.get("name"),
            "email": doc.get("email"),
            "phone": doc.get("phone"),
            "skills": doc.get("skills", []),
            "experience": doc.get("experience", []),
            "education": doc.get("education", []),
            "location": doc.get("location"),
            "total_experience": doc.get("total_experience"),
            "semantic": doc.get("semantic")
        })

    if not results:
        return jsonify({"ok": False, "message": "No data to show"}), 404

    return jsonify({"count": len(results), "results": results})

if __name__ == "__main__":
    app.run(debug=True, port=5000)'''






'''from flask import Flask, request, jsonify
import os, json, re, uuid
from werkzeug.utils import secure_filename
from pdfminer.high_level import extract_text as extract_pdf_text
from dotenv import load_dotenv
import docx
from pymongo import MongoClient
from bson import ObjectId
from flask_cors import CORS
from datetime import datetime
from openai import OpenAI

# ==========================
# CONFIG
# ==========================
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
ALLOWED_EXT = {"pdf", "docx", "txt"}

# Load API key
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
if not os.getenv("OPENAI_API_KEY"):
    raise RuntimeError("OPENAI_API_KEY not set in .env")

app = Flask(__name__)
CORS(app)

# ==========================
# MONGODB INIT
# ==========================
client_mongo = MongoClient("mongodb://localhost:27017/")
db = client_mongo["cv_analyzer"]
candidates_collection = db["candidates"]
candidates_collection.create_index("email", unique=True, sparse=True)

# ==========================
# FILE TEXT EXTRACTION
# ==========================
def extract_text(path):
    ext = path.split(".")[-1].lower()
    if ext == "pdf":
        return extract_pdf_text(path)
    elif ext == "docx":
        doc = docx.Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    elif ext == "txt":
        return open(path, "r", encoding="utf-8", errors="ignore").read()
    return ""

# ==========================
# HELPER: JSON SAFE PARSING
# ==========================
def safe_json_parse(text, fallback=None):
    m = re.search(r"\{(?:.|\n)*\}", text)
    if not m:
        return fallback or {}
    try:
        return json.loads(m.group())
    except Exception as e:
        print("JSON parse error:", e, "RAW:", m.group()[:500])
        return fallback or {}

# ==========================
# OPENAI PARSING
# ==========================
def parse_resume(resume_text):
    prompt = f"""
    Extract the following details from the resume and return only valid JSON.
    If a field is missing, return null or an empty list. Dates for experience 
    should include "start" and "end" in format 'MMM YYYY' if possible.

    {{
      "name": "",
      "email": "",
      "phone": "",
      "location": "",
      "education": [],
      "experience": [
         {{"company": "", "role": "", "start": "", "end": ""}}
      ],
      "skills": []
    }}

    Resume:
    {resume_text[:20000]}
    """

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0
        )
        text = resp.choices[0].message.content
        return safe_json_parse(text, fallback={
            "name": None, "email": None, "phone": None,
            "skills": [], "experience": [], "education": [], "location": None
        })
    except Exception as e:
        print("OpenAI API error:", e)
        return {
            "name": None, "email": None, "phone": None,
            "skills": [], "experience": [], "education": [], "location": None
        }

def semantic_analysis(resume_text):
    prompt = f"""
    Analyze the resume and return only JSON:
    {{
      "summary": "Short 3-4 line professional summary",
      "domain": "IT / Finance / Marketing / Healthcare / Other",
      "role_category": "Developer / Designer / Manager / Analyst / Other",
      "strengths": ["..."],
      "weaknesses": ["..."]
    }}

    Resume:
    {resume_text[:5000]}
    """

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
            temperature=0
        )
        text = resp.choices[0].message.content
        return safe_json_parse(text, fallback={
            "summary": None, "domain": None, "role_category": None,
            "strengths": [], "weaknesses": []
        })
    except Exception as e:
        print("Semantic API error:", e)
        return {
            "summary": None, "domain": None, "role_category": None,
            "strengths": [], "weaknesses": []
        }

# ==========================
# EXPERIENCE CALCULATION
# ==========================
def calculate_total_experience(experiences):
    total_months = 0
    for exp in experiences:
        start_str = exp.get("start")
        end_str = exp.get("end")
        if not start_str:
            continue

        try:
            start = datetime.strptime(start_str, "%b %Y")
        except:
            try:
                start = datetime.strptime(start_str, "%Y")
            except:
                continue

        if end_str:
            try:
                end = datetime.strptime(end_str, "%b %Y")
            except:
                try:
                    end = datetime.strptime(end_str, "%Y")
                except:
                    end = datetime.today()
        else:
            end = datetime.today()

        total_months += (end.year - start.year) * 12 + (end.month - start.month)

    total_years = round(total_months / 12, 2)
    return {
        "years": total_years,
        "formatted": f"{int(total_months/12)} years {total_months%12} months"
    }

# ==========================
# ROUTES
# ==========================
@app.route("/upload", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file"}), 400

    file = request.files["file"]
    filename = secure_filename(file.filename)
    if not filename or filename.split(".")[-1].lower() not in ALLOWED_EXT:
        return jsonify({"error": "Invalid file type"}), 400

    unique_name = f"{uuid.uuid4().hex}_{filename}"
    path = os.path.join(UPLOAD_DIR, unique_name)
    file.save(path)

    text = extract_text(path)
    print("=== Extracted Resume Text ===")
    print(text[:1000])

    parsed = parse_resume(text)
    semantic = semantic_analysis(text)
    parsed["total_experience"] = calculate_total_experience(parsed.get("experience", []))

    # Duplicate check
    existing = None
    if parsed.get("email"):
        existing = candidates_collection.find_one({"email": parsed.get("email")})
    if not existing:
        existing = candidates_collection.find_one({"raw_text": text})

    if existing:
        return jsonify({
            "ok": False,
            "message": "This CV is already in the database",
            "id": str(existing["_id"]),
            "parsed": {
                "name": existing.get("name"),
                "email": existing.get("email"),
                "phone": existing.get("phone"),
                "skills": existing.get("skills"),
                "experience": existing.get("experience"),
                "education": existing.get("education"),
                "location": existing.get("location"),
                "total_experience": existing.get("total_experience"),
                "semantic": existing.get("semantic")
            }
        }), 409

    candidate_doc = {
        "name": parsed.get("name"),
        "email": parsed.get("email"),
        "phone": parsed.get("phone"),
        "skills": parsed.get("skills") or [],
        "experience": parsed.get("experience") or [],
        "education": parsed.get("education") or [],
        "location": parsed.get("location"),
        "total_experience": parsed.get("total_experience"),
        "semantic": semantic,
        "raw_text": text
    }
    inserted = candidates_collection.insert_one(candidate_doc)

    return jsonify({"ok": True, "id": str(inserted.inserted_id), "parsed": parsed, "semantic": semantic})

@app.route("/search")
def search():
    skill_query = request.args.get("skill", "")
    name_query = request.args.get("name", "")
    email_query = request.args.get("email", "")

    query_parts = []
    if skill_query:
        skills = [s.strip() for s in skill_query.split(",")]
        query_parts.append({"$or": [{"skills": {"$regex": s, "$options": "i"}} for s in skills]})
    if name_query:
        query_parts.append({"name": {"$regex": re.escape(name_query), "$options": "i"}})
    if email_query:
        query_parts.append({"email": {"$regex": re.escape(email_query), "$options": "i"}})

    query = {"$and": query_parts} if query_parts else {}
    cursor = candidates_collection.find(query)

    results = []
    for doc in cursor:
        results.append({
            "id": str(doc["_id"]),
            "name": doc.get("name"),
            "email": doc.get("email"),
            "phone": doc.get("phone"),
            "skills": doc.get("skills", []),
            "experience": doc.get("experience", []),
            "education": doc.get("education", []),
            "location": doc.get("location"),
            "total_experience": doc.get("total_experience"),
            "semantic": doc.get("semantic")
        })

    if not results:
        return jsonify({"ok": False, "message": "No data to show"}), 404

    return jsonify({"count": len(results), "results": results})

if __name__ == "__main__":
    app.run(debug=True, port=5000)'''
